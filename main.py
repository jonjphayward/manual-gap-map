import datetime
import json
import math
import os
from docx import Document
import timecode as tc

def is_float(value):
    """Check if a value can be converted to float."""
    try:
        float(value)
        return True
    except ValueError:
        return False

def get_smpte_frames(time, framerate, smpte_timecode):
    # Adjust framerate for drop frame timecodes
    if framerate in [23.976, 23.98]:
        adjusted_fps = 24 * (1000 / 1001)
    elif framerate == 29.97:
        adjusted_fps = 30 * (1000 / 1001)
    elif framerate == 59.94:
        adjusted_fps = 60 * (1000 / 1001)
    else:
        adjusted_fps = framerate

    # Calculate total frames
    total_frames = int(round(time * adjusted_fps))

    # Calculate frames, seconds, minutes, and hours
    frames = total_frames % int(round(adjusted_fps))
    total_seconds = total_frames // int(round(adjusted_fps))
    seconds = total_seconds % 60
    total_minutes = total_seconds // 60
    minutes = total_minutes % 60
    hours = total_minutes // 60

    tc1 = tc.Timecode(framerate, f"{hours:02}:{minutes:02}:{seconds:02}:{frames:02}")

    if smpte_timecode is not None:
        tc2 = tc.Timecode(framerate, f"{smpte_timecode[0]:02}:{smpte_timecode[1]:02}:{smpte_timecode[2]:02}:{smpte_timecode[3]:02}")
        tc3 = tc1 + tc2
        return str(tc3)
    else:
        return f"{hours:02}:{minutes:02}:{seconds:02}:{frames:02}"

def get_smpte_milli(time, smpte_timecode):
    seconds = int(math.floor(time))
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    seconds = seconds % 60
    milliseconds = int(round((time - math.floor(time)) * 1000))

    if smpte_timecode is not None:
        milliseconds += int(smpte_timecode[3])
        extra_seconds = milliseconds // 1000
        milliseconds = milliseconds % 1000
        seconds += extra_seconds + int(smpte_timecode[2])
        extra_minutes = seconds // 60
        seconds = seconds % 60
        minutes += extra_minutes + int(smpte_timecode[1])
        hours += (minutes // 60) + int(smpte_timecode[0])
        minutes = minutes % 60

    return f"{hours:02}:{minutes:02}:{seconds:02},{milliseconds:03}"

def get_gaps(framerate, is_framerate_choice, output_text, output_csv, file_path, smpte_timecode, data):
    """Process gaps in speech based on the given parameters."""
    gap_count = 0
    document = Document()
    paragraph = document.add_paragraph("")

    g_thresh = input("Please type a gap threshold: ")
    if not g_thresh.isdigit():
        raise ValueError("Threshold must be a number.")

    gap_duration = float(g_thresh)
    time = 0.0
    previous_words = []

    # File operations for output
    filename, file_extension = os.path.splitext(file_path)
    file_handlers = open_output_files(output_text, output_csv, filename)

    for key, value in enumerate(data["results"]):
        word = value["alternatives"][0]["content"]
        if len(previous_words) == 5:
            del previous_words[0]
        previous_words.append(word)

        gap = float(value["start_time"]) - time
        if gap > gap_duration:
            gap_count += 1
            start_time = datetime.timedelta(seconds=float(time))
            end_time = datetime.timedelta(seconds=float(value["start_time"]))

            paragraph.add_run(f"\n\n*****Audio gap found*****\nPrevious words: {' '.join(previous_words)}\n")
            if is_framerate_choice:
                paragraph.add_run(f"Start TC: {get_smpte_frames(time, framerate, smpte_timecode)}\n")
                paragraph.add_run(f"End TC: {get_smpte_frames(float(value['start_time']), framerate, smpte_timecode)}\n")
            else:
                paragraph.add_run(f"Start: {get_smpte_milli(time, smpte_timecode)}\n")
                paragraph.add_run(f"End: {get_smpte_milli(float(value['start_time']), smpte_timecode)}\n")
            paragraph.add_run(f"Duration: {gap} seconds")

            if output_text:
                text_file = file_handlers.get("text_file")
                text_file.write(f"{gap_count}\n")
                if is_framerate_choice:
                    text_file.write(f"{get_smpte_frames(time, framerate, smpte_timecode)} --> {get_smpte_frames(float(value['start_time']), framerate, smpte_timecode)}\n")
                else:
                    text_file.write(f"{get_smpte_milli(time, smpte_timecode)} --> {get_smpte_milli(float(value['start_time']), smpte_timecode)}\n")
                text_file.write(f"{' '.join(previous_words)}\n\n")

            if output_csv:
                csv_file = file_handlers.get("csv_file")
                if is_framerate_choice:
                    csv_file.write(f'"{get_smpte_frames(time, framerate, smpte_timecode)}","{get_smpte_frames(float(value["start_time"]), framerate, smpte_timecode)}","{" ".join(previous_words)}",{int(gap*1000)}\n')
                else:
                    csv_file.write(f'"{get_smpte_milli(time, smpte_timecode)}","{get_smpte_milli(float(value["start_time"]), smpte_timecode)}","{" ".join(previous_words)}",{int(gap*1000)}\n')

        time = float(value["start_time"])

    close_output_files(file_handlers)
    return gap_count, document

def open_output_files(output_text, output_csv, file_path):
    """Open file handlers for output files."""
    handlers = {}

    # Extracting the base name of the file for the output file name
    base_filename = os.path.basename(file_path)
    filename, file_extension = os.path.splitext(base_filename)

    if output_text:
        handlers["text_file"] = open(filename + ".srt", "w+", encoding="utf8")
    if output_csv:
        handlers["csv_file"] = open(filename + ".csv", "w+", encoding="utf8")
        handlers["csv_file"].write("Input Timecode,Output Timecode,Dialogue,Max time in ms\n")
    return handlers

def close_output_files(handlers):
    """Close opened file handlers."""
    for handler in handlers.values():
        handler.close()

def get_user_preferences():
    """Get user preferences for timecode and framerate."""
    smpte_timecode = None
    framerate_choice = input("Do you want your results in SMPTE timecode? (Y / N) ").lower() in ["y", "yes"]
    framerate = 0

    if framerate_choice:
        smpte_timecode = get_start_timecode(framerate_choice)
        framerate = get_framerate()
    else:
        # If SMPTE timecode is not chosen, then default to hours:minutes:seconds,milliseconds format
        smpte_timecode = get_start_timecode(framerate_choice)

    return smpte_timecode, framerate, framerate_choice

def get_start_timecode(framerate_choice):
    """Get start timecode from the user."""
    get_start_timecode = input("Do you know the start timecode? (Y / N) ").lower() in ["y", "yes"]
    if get_start_timecode:
        timecode_hours = get_validated_input("Please enter the start timecode's hours value: ", int)
        timecode_mins = get_validated_input("Please enter the start timecode's minutes value: ", int)
        timecode_secs = get_validated_input("Please enter the start timecode's seconds value: ", int)
        if framerate_choice:
            timecode_frames = get_validated_input("Please enter the start timecode's frames value: ", int)
        else:
            timecode_frames = get_validated_input("Please enter the start timecode's milliseconds value: ", int)
            

        return [timecode_hours, timecode_mins, timecode_secs, timecode_frames]
    return None

def get_framerate():
    """Get framerate from the user."""
    while True:
        display_framerate = input("Please type the framerate: ")
        if is_float(display_framerate):
            return float(display_framerate)
        else:
            print("Invalid input. Please type a numeric framerate.")

def get_output_preferences():
    """Get user preferences for output file formats."""
    output_text = input("Do you want to output an srt file? (Y / N) ").lower() in ["y", "yes"]
    output_docx = input("Do you want to output a docx file? (Y / N) ").lower() in ["y", "yes"]
    output_csv = input("Do you want to output a csv file? (Y / N) ").lower() in ["y", "yes"]

    return output_text, output_docx, output_csv

def get_validated_input(prompt, input_type):
    """Get validated input of a specific type from the user."""
    while True:
        user_input = input(prompt)
        try:
            return input_type(user_input)
        except ValueError:
            print(f"Invalid input. Please enter a valid {input_type.__name__}.")

def main():
    """Main function to run the script."""
    input_folder = os.path.join(os.getcwd(), "input")
    if not os.path.exists(input_folder):
        os.makedirs(input_folder)
    if not os.listdir(input_folder):
        print("Please put a file into the input folder")
        return

    for file in os.listdir(input_folder):
        file_path = os.path.join(input_folder, file)
        with open(file_path, encoding="utf8") as f:
            data = json.load(f)

        smpte_timecode, framerate, is_framerate_choice = get_user_preferences()
        output_text, output_docx, output_csv = get_output_preferences()

        gap_count, document = get_gaps(framerate, is_framerate_choice, output_text, output_csv, file_path, smpte_timecode, data)

        print(f"\nTotal number of gaps in the JSON: {gap_count}")
        if output_docx:
            base_filename = os.path.basename(file_path)
            filename = os.path.splitext(base_filename)[0]
            document.save(filename + '.docx')

if __name__ == "__main__":
    main()
    input("Press Enter to exit...")
